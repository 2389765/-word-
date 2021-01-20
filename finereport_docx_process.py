import os
import re
import asyncio
import pythoncom
import win32com.client
import time

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from docxcompose.composer import Composer


word = win32com.client.DispatchEx('Word.Application')
word.visible = False
word.DisplayAlerts = False

root_path = os.getcwd()


class doc_process:

    def __init__(self, text, table, target):
        self.new_doc = Document()
        self.table_path = self.filetype_check(table)
        self.doc = Document(self.filetype_check(text))
        self.to_file = target
        

    def run(self):
        self.extract_text()
        asyncio.run(self.merge_docx())
        asyncio.run(self.modify())
        return True

    def filetype_check(self,path):
        filename, filetype = os.path.splitext(path)
        if filetype == '.doc' or filetype == '.rtf':
            doc = word.Documents.Open(path)
            doc.SaveAs("{}.docx".format(filename), 12)
            print("{}.docx 已生成".format(filename))
            doc.Close()
            return "{}.docx".format(filename)
        elif filetype == '.docx':
            return path    
        
        
    def extract_picture(self):
        """
        按照图片产生的顺序保存图片的_rels，之后通过该值导出图片并回插
        :return: list：按出现顺序排列好的图片rels
        """
        pictures = []
        for rel in self.doc.part._rels:
            if "image" in self.doc.part._rels[rel].target_ref:
                pictures.append((rel, int(re.sub('[^0-9]', ' ', self.doc.part._rels[rel].target_ref))))
        return [_[0] for _ in sorted(pictures, key=lambda x: x[1])]

    def extract_table(self, table):
        """
        筛选出唯一的cell（python.docx中合并的单元格会被视为多个一样的1×1的cell
        :param table: 需要处理的table
        :return: 唯一的cell的index
        """
        row_cells, column_cells = [], []
        index = []
        width, length = len(table.columns), len(table.rows)
        k = 0
        for row in table.rows:
            for cell in row.cells:
                if cell not in row_cells:
                    index.append([k // width, k % width])
                    row_cells.append(cell)
                k += 1
        k = 0
        for column in table.columns:
            for cell in column.cells:
                if cell not in column_cells:
                    column_cells.append(cell)
                elif [k % length, k // length] in index:
                    index.remove([k % length, k // length])
                k += 1
        return index
    
    
    def extract_text(self):
        """
        提取出表格中的文本并且写入到新的文档中
        :return:
        """
        image_index = 0
        pictures = self.extract_picture()
        title_encountered = False
        for table in self.doc.tables:
            index = self.extract_table(table)

            for _ in index:
                if not len(table.rows[_[0]].cells[_[1]].text) == 0:
                    for paragraph in table.rows[_[0]].cells[_[1]].paragraphs:
                        # sub-title
                        if any([True if t in paragraph.text[0] else False for t in
                                ['一', '二', '三', '四', '五', '六', '七', '八', '九']]):
                            para_heading = self.new_doc.add_heading("", level=2)
                            run = para_heading.add_run(paragraph.text)

                        # title
                        elif any([True if t in paragraph.text[-4:] else False for t in ['运行月报', '分析月报']]):
                            para_heading = self.new_doc.add_heading("", level=1)
                            run = para_heading.add_run(paragraph.text)

                        # picture
                        elif '图' in paragraph.text[:2]:
                            para = self.new_doc.add_paragraph()
                            run = para.add_run(paragraph.text)
                            with open(root_path + "\\image.png", "wb") as p:
                                p.write(self.doc.part._rels[pictures[image_index]].target_part.blob)
                                image_index += 1
                            pic = self.new_doc.add_picture(root_path + "\\image.png", width=Inches(5.5))                        
                        
                        # text
                        else:
                            para = self.new_doc.add_paragraph()
                            # 防止标题被识别为正文
                            if len(paragraph.text.split('。', 1)) > 1:
                                run = para.add_run(paragraph.text)


        self.new_doc.add_page_break()
        self.new_doc.save(root_path + '\\temporary.docx')
        return

    async def InsertFile(self,doc_to_insert,from_doc_path):
        doc_to_insert.Application.Selection.InsertFile(from_doc_path)
    
    async def merge_docx(self):
        # time.sleep(30)  # 电脑速度慢时，需要等待以防前面的word修改程序未完成
        # pythoncom.CoInitialize()
        # word = win32com.client.DispatchEx('Word.Application')
        # output = word.Documents.Add()
        # insert1 = asyncio.create_task(self.InsertFile(output,root_path + '\\temporary.docx'))
        # insert2 = asyncio.create_task(self.InsertFile(output,self.table_path))
        # await insert1
        # await insert2
        # # output.Application.Selection.InsertFile(root_path + '\\temporary.docx')
        # # output.Application.Selection.InsertFile(self.table_path)
        # output.SaveAs(self.to_file) 
        # word.Quit()
        
        doc = Document(root_path + '\\temporary.docx')
        cp = Composer(doc)
        cp.append(Document(self.table_path))
        doc.save(self.to_file)
        return
    
    async def font_setting(self,run,text_level):
        font_color = RGBColor(0,0,0)
        
        if text_level == '标题':
            run.font.name = u'黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
            run.font.color.rgb = font_color
            run.font.size = Pt(20)
            run.bold = False
            run.italic = False
            
        elif text_level == '子标题':
            run.font.name = u'黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
            run.font.color.rgb = font_color
            run.font.size = Pt(16)
            run.bold = False
            run.italic = False
        
        elif text_level == '开头':
            run.font.name = u'楷体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
            run.font.color.rgb = font_color
            run.font.size = Pt(16)
            run.bold = True
            run.italic = False
            
        elif text_level == '正文':
            run.font.name = u'仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            run.font.color.rgb = font_color
            run.font.size = Pt(16)
            run.bold = False
            run.italic = False
        
        elif text_level == '图标题':
            run.font.name = u'黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
            run.font.color.rgb = font_color
            run.font.size = Pt(12)
            run.bold = False
            run.italic = False
        return

    
    async def modify(self):
        task_list = []
        # 文档
        doc = Document(self.to_file)

        # 标题
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in doc.paragraphs[0].runs: task_list.append(asyncio.create_task(self.font_setting(run,'标题')))


        # 全文1.5倍行距
        for para in doc.paragraphs[1:]:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            # 图片所在段落
            if len(para.text) == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # 各级小标题
            elif any([True if t in para.text[0] else False for t in ['一', '二', '三', '四', '五', '六', '七', '八', '九']]):
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs: task_list.append(asyncio.create_task(self.font_setting(run,'子标题')))

            # 调节图的语句
            elif '图' in para.text[:2]:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs: task_list.append(asyncio.create_task(self.font_setting(run,'图标题')))               
            # 正文
            elif len(para.text.split('。', 1)) > 1:
                # 调节正文文本
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # 统一每段的缩进为两格
                run = para.runs[0]
                t = re.sub(r"/\s+/", "", run.text)
                run.text = r"  " + re.sub(r"\xa0", "", t)
                
                # text：该para下的所有文本，切分为了开头和正文
                # 使用for循环将该para置空（重新写开头(第一个run)和正文(最后一个run)，解决para的run多于一个的问题）
                text = "".join([_.text for _ in para.runs]).split("。",1)
                for _ in para.runs: _.text = ''
                
                run.text = text[0] + "。"
                task_list.append(asyncio.create_task(self.font_setting(run,'开头')))
                                    
                run = para.add_run(text[1])
                task_list.append(asyncio.create_task(self.font_setting(run,'正文')))
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for each_task in task_list:
            await each_task
        
        doc.save(self.to_file)
        return

if __name__ == '__main__':
    text_path = 'C:/Users/iceberg/Desktop/2.0增值图文宁夏.doc'
    table_path = 'C:/Users/iceberg/Desktop/2.0增值表格宁夏.doc'
    target_path = 'C:/Users/iceberg/Desktop/2.0增值宁夏.docx'
    doc_process(text_path, table_path, target_path).run()


