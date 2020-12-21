import os
import re
import sys
import shutil
import win32com
import win32com.client
import pythoncom
import flask, json
from flask import request, render_template

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt

server = flask.Flask(__name__)
root_path = os.getcwd()
# 保存上传上来的文件
UPLOAD_FOLDER = root_path + '\\UPLOADS'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

class doc_process:

    def __init__(self, text, table, target):
        self.table_path = table
        self.to_file = target
        self.new_doc = Document()
        self.doc = Document(text)

    def run(self):
        self.extract_picture()
        self.extract_text()
        self.merge_docx()
        self.modify()
        return True

    def extract_picture(self):
        """
        按照图片产生的顺序保存图片的_rels，之后通过该值导出图片并回插
        :return: list：按出现顺序排列好的图片rels
        """
        pictures = []
        for rel in self.doc.part._rels:
            if "image" in self.doc.part._rels[rel].target_ref:
                pictures.append((rel, int(re.sub('[^0-9]', ' ', self.doc.part._rels[rel].target_ref))))
        return [_[0] for _ in sorted(pictures, key=lambda x: x[1])]

    def extract_table(self, table):
        """
        筛选出唯一的cell（python.docx中合并的单元格会被视为多个一样的1×1的cell
        :param table: 需要处理的table
        :return: 唯一的cell的index
        """
        row_cells, column_cells = [], []
        index = []
        width, length = len(table.columns), len(table.rows)
        k = 0
        for row in table.rows:
            for cell in row.cells:
                if cell not in row_cells:
                    index.append([k // width, k % width])
                    row_cells.append(cell)
                k += 1
        k = 0
        for column in table.columns:
            for cell in column.cells:
                if cell not in column_cells:
                    column_cells.append(cell)
                elif [k % length, k // length] in index:
                    index.remove([k % length, k // length])
                k += 1
        return index

    def extract_text(self):
        """
        提取出表格中的文本并且写入到新的文档中
        :return:
        """
        image_index = 0
        pictures = self.extract_picture()
        for table in self.doc.tables:
            index = self.extract_table(table)

            for _ in index:
                if not len(table.rows[_[0]].cells[_[1]].text) == 0:
                    for paragraph in table.rows[_[0]].cells[_[1]].paragraphs:
                        # sub-title
                        if any([True if t in paragraph.text[0] else False for t in
                                ['一', '二', '三', '四', '五', '六', '七', '八', '九']]):
                            para_heading = self.new_doc.add_heading("", level=2)
                            run = para_heading.add_run(paragraph.text)

                        # title
                        elif any([True if t in paragraph.text[-4:] else False for t in ['运行月报', '分析月报']]):
                            para_heading = self.new_doc.add_heading("", level=1)
                            run = para_heading.add_run(paragraph.text)

                        # text
                        else:
                            text = paragraph.text.split('。', 1)
                            para = self.new_doc.add_paragraph()
                            if len(text) > 1:
                                run = para.add_run(text)

                            elif '图' in text[0][:2]:
                                run = para.add_run(text[0])
                                with open(root_path + "\\image.png", "wb") as p:
                                    p.write(self.doc.part._rels[pictures[image_index]].target_part.blob)
                                    image_index += 1
                                pic = self.new_doc.add_picture(root_path + "\\image.png", height=Inches(3))

        self.new_doc.add_page_break()
        self.new_doc.save(root_path + '\\temporary.docx')
        return

    def merge_docx(self):
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch('Word.Application')
        word.visible = False
        word.DisplayAlerts = False

        shutil.copy(self.table_path, self.to_file)

        x = word.Documents.Open(root_path + '\\temporary.docx')
        # 复制word的所有内容
        word.Selection.WholeStory()
        word.Selection.Copy()
        # 关闭word
        x.Close()

        to_doc = word.Documents.Open(self.to_file)
        to_doc.Range()
        word.Selection.Paste()
        to_doc.Save()
        to_doc.Close()
        word.Quit()
        return

    def modify(self):
        doc = Document(self.to_file)

        def next_run(para):
            for run in para.runs:
                yield run

        # 标题
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in doc.paragraphs[0].runs:
            run.font.name = u'黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
            run.font.size = Pt(20)

        # 全文1.5倍行距
        for para in doc.paragraphs[1:]:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            # 图片所在段落
            if len(para.text) == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # 各级小标题
            elif any([True if t in para.text[0] else False for t in ['一', '二', '三', '四', '五', '六', '七', '八', '九']]):
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = u'黑体'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                    run.bold = False
                    run.italic = False

            # 正文
            else:
                text = para.text.split('。', 1)
                if '图' in text[0][:2]:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.name = u'黑体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                        run.font.size = Pt(12)
                elif len(text) > 1:
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    nr = next_run(para)
                    # 统一每段的缩进为两格
                    run = next(nr)
                    t = run.text
                    t = re.sub(r"/\s+/", "", t)
                    t = re.sub(r"\xa0", "", t)
                    t = r"  " + t
                    run.text = t
                    # 调节正文的字体字号
                    while True:
                        run = next(nr)
                        run.font.name = u'楷体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
                        run.font.size = Pt(16)
                        run.bold = True
                        if "。" in run.text:
                            break
                    try:
                        while True:
                            run = next(nr)
                            run.font.name = u'仿宋'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
                            run.font.size = Pt(16)
                    except StopIteration:
                        pass
                    finally:
                        del nr
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(self.to_file)
        return

def allowed_file(filename):
    ALLOWED_EXTENSIONS = ['docx']
    # 判断文件的扩展名是否在配置项ALLOWED_EXTENSIONS中
    return '.' in filename and \
           filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS

@server.route('/upload', methods=['get', 'post'])
def upload():
        用于对已经在服务器内的文件进行操作，提供的是文件的地址
    # 获取通过url请求传参的数据
    file = request.values.get('source')
    files = file.split(',')
    # 获取url请求传的密码，明文
    target = request.values.get('tar')
    # 判断用户名、密码都不为空，如果不传用户名、密码则username和pwd为None
    if len(files) != 2:
        resu = {'code': 400, 'message': '需要传递文件和表格的文件位置，用半角的,分隔！'}
        return json.dumps(resu, ensure_ascii=False)
    elif type(target) != str:
        resu = {'code': 400, 'message': '需要存放生成文件的位置'}
        return json.dumps(resu, ensure_ascii=False)
    else:
        target_path, tar_file = os.path.split(target)
        if not os.path.exists(target_path):
            os.makedirs(target_path)

        if '表格' in files[0]:
            doc_process(files[1], files[0], target).run()
        else:
            doc_process(files[0], files[1], target).run()
        resu = {'code': 200, 'message': '转换成功'}
        return json.dumps(resu, ensure_ascii=False)
        # if request.method == 'POST':
    #     # 获取上传过来的文件对象
    #     file = request.files['file']
    #     # 检查文件对象是否存在，且文件名合法
    #     if file and allowed_file(file.filename):
    #         # 去除文件名中不合法的内容
    #         # 将文件保存在本地UPLOAD_FOLDER目录下
    #         file.save(os.path.join(UPLOAD_FOLDER, file.filename))

    #         resu = {'code': 200, 'message': '上传成功'}
    #         return json.dumps(resu, ensure_ascii=False)
    #     else:    # 文件不合法
    #         resu = {'code': 400, 'message': '文件格式有误'}
    #         return json.dumps(resu, ensure_ascii=False)
    # else:  # GET方法
    #     return render_template('upload.html')

if __name__ == '__main__':
    server.run(port=7777, debug=True, host='0.0.0.0')
# http://127.0.0.1:7777/process_doc?source=D:\fineReports\software\2.0%E8%BD%AF%E4%BB%B6\2020\6\%E6%B5%B7%E5%8D%97%E7%9C%81\word\2.0%E9%80%9A%E4%BF%A1%E5%9B%BE%E6%96%87.docx,C:\Users\iceberg\Desktop\%E8%87%AA%E5%8A%A8%E5%8C%96word\%E8%A1%A8%E6%A0%BC.docx&tar=C:\Users\iceberg\Desktop\%E8%87%AA%E5%8A%A8%E5%8C%96word\%E5%A4%84%E7%90%86%E7%BB%93%E6%9E%9C\res.docx
