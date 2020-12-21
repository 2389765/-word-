import os
import re
import sys
import shutil
import win32com
import win32com.client
from collections import Counter

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

word = win32com.client.Dispatch('Word.Application')
word.visible = False
word.DisplayAlerts = False

root_path = os.getcwd()


class doc_process:

    def __init__(self, text, table, target):
        self.table_path = table
        self.to_file = target
        self.new_doc = Document()
        self.doc = Document(text)

    def run(self):
        self.extract_picture()
        self.extract_text()
        self.merge_docx()
        self.modify()
        return True

    def extract_picture(self):
        """
        按照图片产生的顺序保存图片的_rels，之后通过该值导出图片并回插
        :return: list：按出现顺序排列好的图片rels
        """
        pictures = []
        for rel in self.doc.part._rels:
            if "image" in self.doc.part._rels[rel].target_ref:
                pictures.append((rel, int(re.sub('[^0-9]', ' ', self.doc.part._rels[rel].target_ref))))
        return [_[0] for _ in sorted(pictures, key=lambda x: x[1])]

    def extract_table(self, table):
        """
        筛选出唯一的cell（python.docx中合并的单元格会被视为多个一样的1×1的cell
        :param table: 需要处理的table
        :return: 唯一的cell的index
        """
        row_cells, column_cells = [], []
        index = []
        width, length = len(table.columns), len(table.rows)
        k = 0
        for row in table.rows:
            for cell in row.cells:
                if cell not in row_cells:
                    index.append([k // width, k % width])
                    row_cells.append(cell)
                k += 1
        k = 0
        for column in table.columns:
            for cell in column.cells:
                if cell not in column_cells:
                    column_cells.append(cell)
                elif [k % length, k // length] in index:
                    index.remove([k % length, k // length])
                k += 1
        return index

    def extract_text(self):
        """
        提取出表格中的文本并且写入到新的文档中
        :return:
        """
        image_index = 0
        pictures = self.extract_picture()
        for table in self.doc.tables:
            index = self.extract_table(table)

            for _ in index:
                if not len(table.rows[_[0]].cells[_[1]].text) == 0:
                    for paragraph in table.rows[_[0]].cells[_[1]].paragraphs:
                        # sub-title
                        if any([True if t in paragraph.text[0] else False for t in
                                ['一', '二', '三', '四', '五', '六', '七', '八', '九']]):
                            para_heading = self.new_doc.add_heading("", level=2)
                            run = para_heading.add_run(paragraph.text)

                        # title
                        elif any([True if t in paragraph.text[-4:] else False for t in ['运行月报', '分析月报']]):
                            para_heading = self.new_doc.add_heading("", level=1)
                            run = para_heading.add_run(paragraph.text)

                        # text
                        else:
                            text = paragraph.text.split('。', 1)
                            para = self.new_doc.add_paragraph()
                            if len(text) > 1:
                                run = para.add_run(text)

                            elif '图' in text[0][:2]:
                                run = para.add_run(text[0])
                                with open(root_path + "\\image.png", "wb") as p:
                                    p.write(self.doc.part._rels[pictures[image_index]].target_part.blob)
                                    image_index += 1
                                pic = self.new_doc.add_picture(root_path + "\\image.png", height=Inches(3))

        self.new_doc.add_page_break()
        self.new_doc.save(root_path + '\\temporary.docx')
        return

    def merge_docx(self):
        shutil.copy(self.table_path, self.to_file)
        doc = word.Documents.Open(root_path + '\\temporary.docx')
        # 复制word的所有内容
        word.Selection.WholeStory()
        word.Selection.Copy()
        # 关闭word
        doc.Close()

        to_doc = word.Documents.Open(self.to_file)
        to_doc.Range()
        word.Selection.Paste()
        to_doc.Save()
        to_doc.Close()
        word.Quit()
        return

    def modify(self):
        doc = Document(self.to_file)
        color = RGBColor(250,0,0) # 字体颜色

        def next_run(para):
            for run in para.runs:
                yield run

        # 标题
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in doc.paragraphs[0].runs:
            run.font.name = u'黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
            run.font.color.rgb = color
            run.font.size = Pt(20)

        # 全文1.5倍行距
        for para in doc.paragraphs[1:]:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            # 图片所在段落
            if len(para.text) == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # 各级小标题
            elif any([True if t in para.text[0] else False for t in ['一', '二', '三', '四', '五', '六', '七', '八', '九']]):
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = u'黑体'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                    run.font.color.rgb = color
                    run.bold = False
                    run.italic = False

            # 正文
            else:
                text = para.text.split('。', 1)
                if '图' in text[0][:2]:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.name = u'黑体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                        run.font.size = Pt(12)
                elif len(text) > 1:
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    nr = next_run(para)
                    # 统一每段的缩进为两格
                    run = next(nr)
                    t = run.text
                    t = re.sub(r"/\s+/", "", t)
                    t = re.sub(r"\xa0", "", t)
                    t = r"  " + t
                    run.text = t
                    # 调节正文的字体字号
                    while True:
                        run = next(nr)
                        run.font.name = u'楷体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
                        run.font.size = Pt(16)
                        run.font.color.rgb = color
                        run.bold = True
                        if "。" in run.text:
                            break
                    try:
                        while True:
                            run = next(nr)
                            run.font.name = u'仿宋'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
                            run.font.size = Pt(16)
                            run.font.color.rgb = color
                    except StopIteration:
                        pass
                    finally:
                        del nr
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(self.to_file)
        return


def change_to_docx(file_path, filetype):
    doc = word.Documents.Open(file_path + '.' + filetype)
    doc.SaveAs("{}.docx".format(file_path), 12)
    doc.Close()


def main(doc_file, target_file):
    """
    遍历目录，得到文件夹位置
    :return:
    """
    for industry in os.listdir(doc_file):
        I = os.path.join(doc_file, industry)
        if not os.path.isfile(I):

            for year in os.listdir(I):
                Y = os.path.join(I, year)
                if not os.path.isfile(Y):

                    for month in os.listdir(Y):
                        M = os.path.join(Y, month)
                        if not os.path.isfile(M):

                            for province in os.listdir(M):
                                P = os.path.join(M, province)
                                if not os.path.isfile(P):

                                    for doc_type in os.listdir(P):
                                        if doc_type == 'word':
                                            D = os.path.join(P, doc_type)

                                            docs = os.listdir(D)
                                            doc_names = []
                                            for file in docs:
                                                filename, filetype = os.path.splitext(file)
                                                if filetype == '.doc' or filetype == '.rtf':
                                                    change_to_docx(os.path.join(D, filename), filetype)
                                                    doc_names.append(filename[:len(filename) - 2])
                                                elif filetype == '.docx':
                                                    doc_names.append(filename[:len(filename) - 2])
                                                else:
                                                    continue
                                            x = Counter(doc_names).most_common()
                                            for each in x:
                                                if each[1] == 2:
                                                    text_path = os.path.join(D, each[0] + '图文.docx')
                                                    table_path = os.path.join(D, each[0] + '表格.docx')
                                                    target_path = target_file + '/%s/%s/%s/%s/%s' % (
                                                    industry, year, month, province, doc_type)
                                                    if not os.path.exists(target_path):
                                                        os.makedirs(target_path)
                                                    doc_process(text_path, table_path, target_path+'/%s.docx'% each[0]).run()

if __name__ == '__main__':
    if len(sys.argv) == 1:
        doc_file = r'D:/fineReports/software'
        target_file = r'D:/fineReports/processed'
        main(doc_file, target_file)
    elif len(sys.argv) == 4:
        text_path = sys.argv[1]
        table_path = sys.argv[2]
        target_path = sys.argv[3]
        doc_process(text_path, table_path, target_path).run()
    else:
        print("命令行参数有误，请输入为：图文文件路径，表格文件路径，输出文件路径")
